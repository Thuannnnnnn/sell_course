import httpx
import json
import tempfile
import os
from typing import Tuple
from docx import Document
from urllib.parse import urlparse
import asyncio

class FileProcessor:
    """Service to process files from URLs and extract text content"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.json']
    
    async def process_url(self, url: str) -> Tuple[str, str]:
        """
        Process a URL and extract text content
        
        Args:
            url: URL to the file
            
        Returns:
            Tuple of (text_content, filename)
        """
        # Download file
        file_content, filename = await self._download_file(url)
        
        # Determine file type and process accordingly
        file_extension = self._get_file_extension(filename)
        
        if file_extension == '.docx':
            text_content = await self._process_docx(file_content)
        elif file_extension == '.json':
            text_content = await self._process_json(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return text_content, filename
    
    async def _download_file(self, url: str) -> Tuple[bytes, str]:
        """Download file from URL"""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                # Extract filename from URL
                parsed_url = urlparse(url)
                filename = os.path.basename(parsed_url.path)
                
                # If no filename in URL, try to get from Content-Disposition header
                if not filename or '.' not in filename:
                    content_disposition = response.headers.get('content-disposition', '')
                    if 'filename=' in content_disposition:
                        filename = content_disposition.split('filename=')[1].strip('"')
                    else:
                        # Default filename based on content type
                        content_type = response.headers.get('content-type', '')
                        if 'json' in content_type:
                            filename = 'downloaded_file.json'
                        elif 'word' in content_type or 'docx' in content_type:
                            filename = 'downloaded_file.docx'
                        else:
                            filename = 'downloaded_file'
                
                return response.content, filename
                
        except httpx.RequestError as e:
            raise Exception(f"Error downloading file from {url}: {str(e)}")
        except httpx.HTTPStatusError as e:
            raise Exception(f"HTTP error {e.response.status_code} when downloading {url}")
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        return os.path.splitext(filename.lower())[1]
    
    async def _process_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Process DOCX file
                doc = Document(temp_file_path)
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(' | '.join(row_text))
                
                return '\n\n'.join(text_content)
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            raise Exception(f"Error processing DOCX file: {str(e)}")
    
    async def _process_json(self, file_content: bytes) -> str:
        """Extract text from JSON file"""
        try:
            # Parse JSON content
            json_data = json.loads(file_content.decode('utf-8'))
            
            # Extract text content from JSON
            text_content = self._extract_text_from_json(json_data)
            
            return text_content
            
        except json.JSONDecodeError as e:
            raise Exception(f"Error parsing JSON file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing JSON file: {str(e)}")
    
    def _extract_text_from_json(self, data, prefix="") -> str:
        """Recursively extract text content from JSON data"""
        text_parts = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                current_prefix = f"{prefix}.{key}" if prefix else key
                
                if isinstance(value, (dict, list)):
                    text_parts.append(self._extract_text_from_json(value, current_prefix))
                elif isinstance(value, str) and value.strip():
                    text_parts.append(f"{current_prefix}: {value.strip()}")
                elif value is not None:
                    text_parts.append(f"{current_prefix}: {str(value)}")
                    
        elif isinstance(data, list):
            for i, item in enumerate(data):
                current_prefix = f"{prefix}[{i}]" if prefix else f"item_{i}"
                
                if isinstance(item, (dict, list)):
                    text_parts.append(self._extract_text_from_json(item, current_prefix))
                elif isinstance(item, str) and item.strip():
                    text_parts.append(f"{current_prefix}: {item.strip()}")
                elif item is not None:
                    text_parts.append(f"{current_prefix}: {str(item)}")
                    
        elif isinstance(data, str) and data.strip():
            text_parts.append(data.strip())
        elif data is not None:
            text_parts.append(str(data))
        
        return '\n'.join(filter(None, text_parts))

